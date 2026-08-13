#!/usr/bin/env python3
"""Markdown仕様書 → 納品用 docx 変換スクリプト

✅️ Ex生産情報一覧1_仕様書.docx のデザイン（配色・表スタイル・章立て）を踏襲しつつ、
クライアント納品文書として必要な要素を追加した docx を生成する。

追加要素:
  - 表紙（機密区分ラベル / 文書管理情報表）
  - 改訂履歴ページ・取扱い注意
  - 目次（静的な見出し一覧。TOC フィールドは使わない）
  - ヘッダー / フッター（文書名・機密区分・文書番号・ページ番号）
  - Word 標準見出しスタイル（ナビゲーションウィンドウ / PDF しおり対応）
  - 表の固定列幅・見出し行のページ跨ぎ繰り返し・ゼブラ

使い方:
  python3 scripts/md_to_docx.py docs/Ex１号機作業指図書_仕様書.md \
      --out "納品/Ex１号機作業指図書_仕様書.docx" \
      --issuer "株式会社〇〇" --client "〇〇株式会社" --rev 1.0 --date 2026-07-27
"""

import argparse
import re
import unicodedata
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor, Twips

# ── フォント ───────────────────────────────────────────────
# 半角カタカナを含むため、欧文=Arial / 和文=BIZ UDPGothic を明示指定する
# （和文フォントを指定しないと環境により Gungsuh へフォールバックし \ が ₩ 表示になる）
FONT_LATIN = "Arial"
FONT_EA = "BIZ UDPGothic"
FONT_CODE_LATIN = "Consolas"
FONT_CODE_EA = "BIZ UDGothic"

# ── 配色（参照 docx 準拠）──────────────────────────────────
COLOR_H1 = RGBColor(0x1F, 0x38, 0x64)
COLOR_H2 = RGBColor(0x2E, 0x5F, 0xA3)
COLOR_H3 = RGBColor(0x2E, 0x75, 0xB6)
COLOR_H4 = RGBColor(0x55, 0x55, 0x55)
COLOR_NOTE = RGBColor(0x3F, 0x3F, 0x3F)
COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)
COLOR_CODE = RGBColor(0x33, 0x33, 0x33)
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_ACCENT = "2e5fa3"
HEX_ACCENT_DARK = "1f3864"
HEX_ROW_ALT = "ebf3fb"
HEX_ROW_BASE = "ffffff"
HEX_CELL_BORDER = "b8cce4"
HEX_CODE_BG = "f2f2f2"
HEX_CODE_BORDER = "aaaaaa"
HEX_NOTE_BG = "fdf6e3"
HEX_NOTE_BORDER = "e0b64a"
HEX_RULE = "cccccc"

# ── サイズ ─────────────────────────────────────────────────
SIZE_TITLE = Pt(30)
SIZE_SUBTITLE = Pt(20)
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(14)
SIZE_H3 = Pt(12)
SIZE_H4 = Pt(11)
SIZE_BODY = Pt(10)
SIZE_TABLE = Pt(8)
SIZE_NOTE = Pt(9)
SIZE_CODE = Pt(9)
SIZE_CHROME = Pt(8)

# ── ページ（A4 縦）─────────────────────────────────────────
PAGE_WIDTH = Twips(11906)
PAGE_HEIGHT = Twips(16838)
MARGIN_X = Twips(850)
MARGIN_Y = Twips(1000)
CONTENT_WIDTH = 11906 - 850 * 2  # = 10206 twips

MIN_COL_WIDTH = 620
NARROW_HEADERS = {"✓", "#", "No.", "No", "版", "章"}
NARROW_COL_WIDTH = 460


# ── 低レベルヘルパ ─────────────────────────────────────────
# Word は各プロパティ要素の出現順に厳格なため、スキーマ順を保って挿入する

ORDER = {
    "w:rPr": ["w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
              "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
              "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
              "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern",
              "w:position", "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect",
              "w:bdr", "w:shd", "w:fitText", "w:vertAlign", "w:rtl", "w:cs",
              "w:em", "w:lang", "w:eastAsianLayout", "w:specVanish",
              "w:oMath", "w:rPrChange"],
    "w:pPr": ["w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
              "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
              "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
              "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
              "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
              "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
              "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
              "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
              "w:rPr", "w:sectPr", "w:pPrChange"],
    "w:tblPr": ["w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
                "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW",
                "w:jc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd",
                "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption",
                "w:tblDescription", "w:tblPrChange"],
    "w:trPr": ["w:cnfStyle", "w:divId", "w:gridBefore", "w:gridAfter",
               "w:wBefore", "w:wAfter", "w:cantSplit", "w:trHeight",
               "w:tblHeader", "w:tblCellSpacing", "w:jc", "w:hidden"],
    "w:tcPr": ["w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge",
               "w:tcBorders", "w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
               "w:tcFitText", "w:vAlign", "w:hideMark"],
}


def _insert(parent, xml):
    """親要素のスキーマ順を守って子要素を差し替え挿入する"""
    new = parse_xml(xml)
    order = ORDER["w:" + parent.tag.split('}')[-1]]
    tag = "w:" + new.tag.split('}')[-1]
    for existing in parent.findall(qn(tag)):
        parent.remove(existing)
    rank = order.index(tag)
    for child in parent:
        child_tag = "w:" + child.tag.split('}')[-1]
        if child_tag in order and order.index(child_tag) > rank:
            child.addprevious(new)
            return new
    parent.append(new)
    return new


def _set_font(run, latin=FONT_LATIN, ea=FONT_EA, size=SIZE_BODY,
              color=COLOR_BLACK, bold=False, italic=False):
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    _insert(run._element.get_or_add_rPr(),
            f'<w:rFonts {nsdecls("w")} w:ascii="{latin}" w:hAnsi="{latin}" '
            f'w:cs="{latin}" w:eastAsia="{ea}"/>')


def _shade(element, hex_color):
    """段落 / セルに背景色を設定"""
    pr = element.get_or_add_pPr() if element.tag == qn('w:p') else element.get_or_add_tcPr()
    _insert(pr, f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{hex_color}"/>')


def _para_border(paragraph, side, hex_color, sz=6, space=4):
    _insert(paragraph._element.get_or_add_pPr(),
            f'<w:pBdr {nsdecls("w")}><w:{side} w:val="single" w:sz="{sz}" '
            f'w:space="{space}" w:color="{hex_color}"/></w:pBdr>')


def _table_borders_xml(outer="808080", inner=HEX_CELL_BORDER):
    return f"""<w:tblBorders {nsdecls("w")}>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="{outer}"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="{outer}"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{outer}"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="{outer}"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{inner}"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{inner}"/>
    </w:tblBorders>"""


def _apply_table_layout(table, widths):
    """テーブル幅を固定レイアウトで確定させる"""
    tbl = table._element
    table.autofit = False
    tblPr = tbl.tblPr
    _insert(tblPr, f'<w:tblW {nsdecls("w")} w:w="{sum(widths)}" w:type="dxa"/>')
    _insert(tblPr, _table_borders_xml())
    _insert(tblPr, f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    _insert(tblPr, f'<w:tblLook {nsdecls("w")} w:val="0000"/>')
    for gridCol, w in zip(tbl.find(qn('w:tblGrid')).findall(qn('w:gridCol')), widths):
        gridCol.set(qn('w:w'), str(w))


def _style_cell(cell, width, bg, margin=(60, 100)):
    cell.width = Twips(width)
    tcPr = cell._element.get_or_add_tcPr()
    _insert(tcPr, f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{bg}"/>')
    _insert(tcPr, f'<w:tcMar {nsdecls("w")}>'
                  f'<w:top w:w="{margin[0]}" w:type="dxa"/>'
                  f'<w:left w:w="{margin[1]}" w:type="dxa"/>'
                  f'<w:bottom w:w="{margin[0]}" w:type="dxa"/>'
                  f'<w:right w:w="{margin[1]}" w:type="dxa"/></w:tcMar>')
    _insert(tcPr, f'<w:vAlign {nsdecls("w")} w:val="top"/>')


def _page_break_before(paragraph):
    _insert(paragraph._element.get_or_add_pPr(),
            f'<w:pageBreakBefore {nsdecls("w")} w:val="true"/>')


def _field(paragraph, instr, size=SIZE_CHROME, color=COLOR_MUTED):
    """PAGE / NUMPAGES 等のフィールドを挿入（未更新時の表示値付き）"""
    def add(xml=None, text=None):
        run = paragraph.add_run()
        _set_font(run, size=size, color=color)
        if xml is not None:
            run._element.append(parse_xml(xml))
        if text is not None:
            run.text = text

    add(xml=f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    add(xml=f'<w:instrText {nsdecls("w")} xml:space="preserve">{instr}</w:instrText>')
    add(xml=f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    add(text="1")
    add(xml=f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')


def _display_width(text):
    """全角を 2、半角を 1 として表示幅を数える"""
    return sum(2 if unicodedata.east_asian_width(c) in "WFA" else 1 for c in text)


# ── インライン Markdown ────────────────────────────────────

def parse_inline(text):
    """**太字** / `コード` / [text](url) を (text, bold, code) 列に分解"""
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    parts = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(2):
            parts.append((m.group(2), True, False))
        else:
            parts.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts or [(text, False, False)]


def plain_text(text):
    return "".join(t for t, _, _ in parse_inline(text))


def add_rich_paragraph(doc_or_cell, text, size=SIZE_BODY, color=COLOR_BLACK,
                       bold=False, alignment=None, space_before=None,
                       space_after=None, indent_left=None, keep_next=False,
                       paragraph=None):
    p = paragraph if paragraph is not None else doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if indent_left is not None:
        pf.left_indent = indent_left
    if keep_next:
        pf.keep_with_next = True

    for li, line in enumerate(text.split("\n")):
        if li:
            br = p.add_run()
            _set_font(br, size=size, color=color)
            br._element.append(parse_xml(f'<w:br {nsdecls("w")}/>'))
        for chunk, is_bold, is_code in parse_inline(line):
            run = p.add_run(chunk)
            if is_code:
                _set_font(run, FONT_CODE_LATIN, FONT_CODE_EA, size, COLOR_CODE,
                          bold or is_bold)
            else:
                _set_font(run, size=size, color=color, bold=bold or is_bold)
    return p


# ── 表 ─────────────────────────────────────────────────────

def _column_widths(header, rows):
    """列内容の表示幅から固定列幅（twips）を決める

    見出しが語中で折り返さない幅を各列の下限とし、残り幅を内容量で按分する。
    """
    ncols = len(header)
    weights, minima = [], []
    for ci in range(ncols):
        head = plain_text(header[ci]).strip()
        cells = [plain_text(r[ci]) if ci < len(r) else "" for r in rows]
        # セル内改行は行ごとに折り返すため、最長行の幅で評価する
        widths = sorted(max((_display_width(l) for l in c.split("\n")), default=0)
                        for c in cells) or [0]
        # 極端に長い 1 セルに引きずられないよう 85 パーセンタイルを採用
        idx = min(len(widths) - 1, int(len(widths) * 0.85))
        weights.append(max(_display_width(head), widths[idx], 2))
        if head in NARROW_HEADERS:
            minima.append(NARROW_COL_WIDTH)
        else:
            # 8pt の 1 文字 ≒ 92 twips + セル余白。
            # 見出しと、短い内容（Event / Sub 等）が折り返さない幅を下限とする
            need = max(_display_width(head), min(widths[-1], 10))
            minima.append(min(2400, max(MIN_COL_WIDTH, need * 92 + 260)))

    if sum(minima) > CONTENT_WIDTH:            # 下限の合計が本文幅を超える場合は縮尺
        scale = CONTENT_WIDTH / sum(minima)
        minima = [int(m * scale) for m in minima]

    widths = {}
    pool = list(range(ncols))
    budget = CONTENT_WIDTH
    while pool:
        total = sum(weights[ci] for ci in pool) or 1
        undersized = [ci for ci in pool
                      if budget * weights[ci] / total < minima[ci]]
        if not undersized:
            for ci in pool:
                widths[ci] = int(budget * weights[ci] / total)
            break
        for ci in undersized:
            widths[ci] = minima[ci]
            budget -= minima[ci]
            pool.remove(ci)
        if budget <= 0:  # 下限の合計が幅を超える極端なケース
            for ci in pool:
                widths[ci] = minima[ci]
            break

    # 丸め誤差を最も広い列で吸収し、合計を必ず本文幅に一致させる
    widest = max(range(ncols), key=lambda ci: widths[ci])
    widths[widest] += CONTENT_WIDTH - sum(widths.values())
    return [widths[ci] for ci in range(ncols)]


def _write_cell(cell, text, width, bg, size=SIZE_TABLE, bold=False,
                color=COLOR_BLACK, margin=(70, 110), center=False):
    _style_cell(cell, width, bg, margin)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    add_rich_paragraph(None, text, size=size, color=color, bold=bold,
                       paragraph=p,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER if center else None)


def add_table(doc, header, rows):
    ncols = len(header)
    widths = _column_widths(header, rows)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _apply_table_layout(table, widths)

    centered = [plain_text(h).strip() in NARROW_HEADERS for h in header]

    # 見出し行（ページを跨いだら繰り返す）
    hdr_row = table.rows[0]
    _insert(hdr_row._element.get_or_add_trPr(),
            f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
    for ci, text in enumerate(header):
        _write_cell(hdr_row.cells[ci], text.strip(), widths[ci], HEX_ACCENT,
                    bold=True, color=COLOR_WHITE, center=centered[ci])

    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        _insert(tr._element.get_or_add_trPr(),
                f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        bg = HEX_ROW_ALT if ri % 2 == 0 else HEX_ROW_BASE
        for ci in range(ncols):
            _write_cell(tr.cells[ci], row[ci].strip() if ci < len(row) else "",
                        widths[ci], bg, center=centered[ci])
    return table


def add_code_block(doc, lines):
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.left_indent = Twips(240)
        _shade(p._element, HEX_CODE_BG)
        _para_border(p, "left", HEX_CODE_BORDER, sz=8)
        run = p.add_run(line)
        _set_font(run, FONT_CODE_LATIN, FONT_CODE_EA, SIZE_CODE, COLOR_CODE)


def add_note(doc, text):
    """blockquote → コールアウト"""
    p = add_rich_paragraph(doc, text, size=SIZE_NOTE, color=COLOR_NOTE,
                           indent_left=Twips(120),
                           space_before=Pt(4), space_after=Pt(6))
    _shade(p._element, HEX_NOTE_BG)
    _para_border(p, "left", HEX_NOTE_BORDER, sz=12)
    return p


# ── Markdown パーサ ───────────────────────────────────────

def parse_md(md_path):
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    tokens = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            code, i = [], i + 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            tokens.append(("code", code))
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            tokens.append(("heading", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        if line.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            if len(block) >= 2:
                header = [c.strip() for c in block[0].split("|")[1:-1]]
                rows = [[c.strip() for c in tl.split("|")[1:-1]] for tl in block[2:]]
                tokens.append(("table", header, rows))
            continue

        if line.startswith("> ") or line.strip() == ">":
            bq = []
            while i < len(lines) and (lines[i].startswith("> ") or lines[i].strip() == ">"):
                bq.append(lines[i][2:] if lines[i].startswith("> ") else "")
                i += 1
            tokens.append(("blockquote", bq))
            continue

        if re.match(r'^---+\s*$', line.strip()):
            tokens.append(("hr",))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        m = re.match(r'^(\s*)(?:[-*]|\d+\.)\s+(.*)', line)
        if m:
            tokens.append(("list_item", m.group(2).strip(), len(m.group(1))))
            i += 1
            continue

        # 空行までの連続行は 1 段落として扱う（Markdown の段落規則）
        buf = []
        while i < len(lines) and lines[i].strip() \
                and not re.match(r'^(#{1,6}\s|```|\||>\s|>$|---+\s*$)', lines[i]) \
                and not re.match(r'^\s*(?:[-*]|\d+\.)\s', lines[i]):
            buf.append(lines[i].strip())
            i += 1
        tokens.append(("paragraph", " ".join(buf)))
    return tokens


# ── 再構築パス（箇条書き主体の記述を表形式へ正規化）──────────

PROC_RE = re.compile(
    r'^(?P<name>.+?)\s*[—–]\s*(?P<kind>Sub|Event|Function|Property)\b(?P<rest>.*)$')
MAX_LABEL = 30


def _mask(text):
    """インラインコード（`...`）と括弧内を伏せた文字列を返す

    「ラベル: 値」の区切り位置を探すとき、`Test()` や （注記：…） の中の
    コロン・括弧を誤検出しないためのマスク処理。
    """
    out = list(text)
    for m in re.finditer(r'`[^`]*`', text):
        for i in range(*m.span()):
            out[i] = "\x00"
    depth = 0
    for i, ch in enumerate(text):
        if out[i] == "\x00":
            continue
        if ch in "（(【「":
            depth += 1
            out[i] = "\x01"
        elif ch in "）)】」" and depth > 0:
            depth -= 1
            out[i] = "\x01"
        elif depth > 0:
            out[i] = "\x01"
    return "".join(out)


def _bullet_groups(items):
    """(インデント, テキスト) 列を「親 → 子リスト」に畳む"""
    groups = []
    for level, text in items:
        if level == 0 or not groups:
            groups.append([text, []])
        else:
            groups[-1][1].append(("　" * (level // 2 - 1)) + "－ " + text)
    return groups


def _split_label(text):
    """「ラベル: 値」「ラベル（値）」を (ラベル, 値) に分解。該当なしは (None, text)"""
    text = text.strip()
    masked = _mask(text)

    for i, ch in enumerate(masked[:MAX_LABEL + 1]):
        if ch in ":：":
            label, value = text[:i].strip(), text[i + 1:].strip()
            if label:
                return label, value        # 末尾コロン（子要素の見出し）は value=""
            break

    # 「ラベル（値）」形式: 末尾が丸ごと括弧で、その前にラベルがある場合のみ
    if masked.endswith("\x01"):
        start = len(masked.rstrip("\x01"))   # 末尾括弧の開き位置
        head = masked[:start]
        if 0 < start <= MAX_LABEL and "\x01" not in head and "\x00" not in head:
            return text[:start].strip(), text[start + 1:-1].strip()
    return None, text


SEG_KV = re.compile(r'^(?P<k>[^（(]{1,20})[（(](?P<v>[^）)]{1,40})[）)]$')


def _split_top_level(text, seps="、,"):
    """括弧やインラインコードの外側にある区切り文字だけで分割する"""
    masked = _mask(text)
    parts, start = [], 0
    for i, ch in enumerate(masked):
        if ch in seps:
            parts.append(text[start:i])
            start = i + 1
    parts.append(text[start:])
    return [p.strip() for p in parts if p.strip()]


def _expand_enumeration(text):
    """「A（説明）、B（説明）」の列挙を複数行に展開する。該当なしは None"""
    parts = _split_top_level(text)
    if len(parts) < 2:
        return None
    rows = []
    for part in parts:
        m = SEG_KV.match(part)
        if not m:
            return None
        rows.append([m.group("k").strip(), m.group("v").strip()])
    return rows


def _bullets_to_token(items):
    """箇条書きブロックを表 / 段落 / 箇条書きのいずれかに変換する"""
    groups = _bullet_groups(items)

    if len(groups) == 1 and not groups[0][1]:
        return ("paragraph", groups[0][0])

    has_child = any(children for _, children in groups)
    rows, labelled = [], 0
    for top, children in groups:
        label, value = _split_label(top)
        if label is None and not children:
            expanded = _expand_enumeration(top)
            if expanded:
                rows.extend(expanded)
                labelled += 1
                continue
        if label is not None:
            labelled += 1
        # ラベルが取れない行は空欄セルにせず「補足」として扱う
        rows.append([label or "補足", "\n".join(filter(None, [value] + children))])

    # 項目名が取れない行が混ざると空欄セルになるため、
    # ほぼ全行がラベル付きのときだけ「項目 / 内容」の 2 列表にする
    if labelled / len(groups) >= 0.8 \
            and sum(1 for k, v in rows if not k or not v) / len(rows) < 0.25:
        return ("table", ["項目", "内容"], rows)

    flat = ["\n".join(filter(None, [top] + children)) for top, children in groups]
    return ("table", ["No", "内容"],
            [[str(i + 1), t] for i, t in enumerate(flat)])


def _fold_procedures(blocks):
    """プロシージャ見出し + 説明を 1 つのモジュール表にまとめる"""
    rows, has_kind = [], False
    for head, body in blocks:
        m = PROC_RE.match(head)
        name = head
        kind = ""
        if m:
            has_kind = True
            name = (m.group("name").strip() + " " + m.group("rest").strip()).strip()
            kind = m.group("kind")
        lines = []
        for tok in body:
            if tok[0] == "paragraph":
                lines.append(tok[1])
            else:                              # list_items
                for level, text in tok[1]:
                    prefix = "・" if level == 0 else ("　" * (level // 2 - 1)) + "－ "
                    lines.append(prefix + text)
        rows.append([name, kind, "\n".join(lines)])

    if has_kind:
        return ("table", ["プロシージャ", "種別", "処理内容"], rows)
    return ("table", ["項目", "内容"], [[r[0], r[2]] for r in rows])


def restructure(tokens):
    """箇条書き主体の記述を表形式に正規化したトークン列を返す"""
    out = []
    i = 0
    while i < len(tokens):
        tok = tokens[i]

        # ① プロシージャ見出し（#### xxx — Sub）の連続を 1 表に集約
        if tok[0] == "heading" and tok[1] == 4 and PROC_RE.match(tok[2]):
            blocks, j = [], i
            while j < len(tokens) and tokens[j][0] == "heading" \
                    and tokens[j][1] == 4 and PROC_RE.match(tokens[j][2]):
                head, body, j = tokens[j][2], [], j + 1
                while j < len(tokens) and tokens[j][0] in ("paragraph", "list_item"):
                    if tokens[j][0] == "list_item":
                        items = []
                        while j < len(tokens) and tokens[j][0] == "list_item":
                            items.append((tokens[j][2], tokens[j][1]))
                            j += 1
                        body.append(("list_item", items))
                    else:
                        body.append(tokens[j])
                        j += 1
                if j < len(tokens) and tokens[j][0] in ("table", "code"):
                    # 表やコードを含むプロシージャは畳まず、個別に見出しを残す
                    if blocks:
                        out.append(_fold_procedures(blocks))
                        blocks = []
                    out.append(("heading", 4, head))
                    out.extend(t if t[0] == "paragraph"
                               else _bullets_to_token(t[1]) for t in body)
                    break
                blocks.append((head, body))
            if blocks:
                out.append(_fold_procedures(blocks))
            i = j
            continue

        # ② 箇条書きブロックを表 / 段落へ
        if tok[0] == "list_item":
            items, j = [], i
            while j < len(tokens) and tokens[j][0] == "list_item":
                items.append((tokens[j][2], tokens[j][1]))
                j += 1
            out.append(_bullets_to_token(items))
            i = j
            continue

        out.append(tok)
        i += 1

    # list_block / list を list_item 列に戻す
    flat = []
    for tok in out:
        if tok[0] in ("list", "list_block"):
            for level, text in tok[1]:
                flat.append(("list_item", text, level))
        else:
            flat.append(tok)
    return flat


META_KEY_ALIASES = {
    "VBA プロジェクトサイズ": "VBA プロジェクト",
    "外部連携ファイル": "外部連携",
    "対象ファイル": "対象ファイル",
}


def extract_meta(tokens):
    title = next((t[2] for t in tokens if t[0] == "heading" and t[1] == 1), "")
    meta = {}
    for tok in tokens:
        if tok[0] == "blockquote":
            for line in tok[1]:
                m = re.match(r'\*\*(.+?)\*\*:\s*(.*)', line)
                if m:
                    key = m.group(1).strip()
                    key = META_KEY_ALIASES.get(key, key)
                    val = plain_text(m.group(2).strip())
                    if key in meta and meta[key]:
                        if val and val not in meta[key]:
                            meta[key] = meta[key] + " / " + val
                    else:
                        meta[key] = val
            if meta:
                break
    return title, meta


# ── 文書スタイル定義 ──────────────────────────────────────

def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.size = SIZE_BODY
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15
    normal.element.get_or_add_rPr().append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_LATIN}" w:hAnsi="{FONT_LATIN}" '
        f'w:cs="{FONT_LATIN}" w:eastAsia="{FONT_EA}"/>'))

    spec = [
        ("Heading 1", SIZE_H1, COLOR_H1, Pt(0), Pt(10)),
        ("Heading 2", SIZE_H2, COLOR_H2, Pt(14), Pt(6)),
        ("Heading 3", SIZE_H3, COLOR_H3, Pt(11), Pt(4)),
        ("Heading 4", SIZE_H4, COLOR_H4, Pt(9), Pt(3)),
    ]
    for name, size, color, before, after in spec:
        st = doc.styles[name]
        st.font.size = size
        st.font.bold = True
        st.font.color.rgb = color
        st.font.name = FONT_LATIN
        st.element.get_or_add_rPr().append(parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_LATIN}" w:hAnsi="{FONT_LATIN}" '
            f'w:cs="{FONT_LATIN}" w:eastAsia="{FONT_EA}"/>'))
        pf = st.paragraph_format
        pf.space_before = before
        pf.space_after = after
        pf.keep_with_next = True
        pf.keep_together = True
        if name == "Heading 1":   # 章の区切りを明確にする下罫線
            _insert(st.element.get_or_add_pPr(),
                    f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" '
                    f'w:space="3" w:color="{HEX_ACCENT}"/></w:pBdr>')

    # 目次スタイル（TOC フィールド更新後もフォントを維持）
    for name, size, color, indent in (("toc 1", Pt(11), COLOR_H1, 0),
                                      ("toc 2", Pt(10), COLOR_H2, 360),
                                      ("toc 3", Pt(9), COLOR_H3, 720)):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.size = size
        st.font.color.rgb = color
        st.element.get_or_add_rPr().append(parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_LATIN}" w:hAnsi="{FONT_LATIN}" '
            f'w:cs="{FONT_LATIN}" w:eastAsia="{FONT_EA}"/>'))
        st.paragraph_format.left_indent = Twips(indent)
        st.paragraph_format.space_after = Pt(2)


def setup_page(doc, doc_title, confidential, footer_left):
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = section.right_margin = MARGIN_X
    section.top_margin = section.bottom_margin = MARGIN_Y
    section.header_distance = Twips(560)
    section.footer_distance = Twips(560)
    section.different_first_page_header_footer = True  # 表紙はヘッダー/フッターなし

    tab = f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="{CONTENT_WIDTH}"/></w:tabs>'

    hp = section.header.paragraphs[0]
    hp._element.get_or_add_pPr().append(parse_xml(tab))
    _para_border(hp, "bottom", HEX_RULE, sz=4)
    run = hp.add_run(doc_title)
    _set_font(run, size=SIZE_CHROME, color=COLOR_MUTED)
    run = hp.add_run("\t" + confidential)
    _set_font(run, size=SIZE_CHROME, color=COLOR_MUTED, italic=True)

    fp = section.footer.paragraphs[0]
    fp._element.get_or_add_pPr().append(parse_xml(tab))
    _para_border(fp, "top", HEX_RULE, sz=4)
    run = fp.add_run(footer_left + "\t")
    _set_font(run, size=SIZE_CHROME, color=COLOR_MUTED)
    _field(fp, "PAGE")
    run = fp.add_run(" / ")
    _set_font(run, size=SIZE_CHROME, color=COLOR_MUTED)
    _field(fp, "NUMPAGES")
    run = fp.add_run(" ページ")
    _set_font(run, size=SIZE_CHROME, color=COLOR_MUTED)

    # 目次は TOC フィールドにしない（開いた直後の更新で頁がすべて 1 になるため）。
    # フッターの PAGE / NUMPAGES は表示時に Word が組版するため、自動更新は付けない。


# ── 表紙・改訂履歴・目次 ──────────────────────────────────

def add_info_table(doc, rows_data, label_width=2400):
    val_width = CONTENT_WIDTH - label_width
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _apply_table_layout(table, [label_width, val_width])

    for ri, (key, val) in enumerate(rows_data):
        _write_cell(table.rows[ri].cells[0], key, label_width, HEX_ROW_ALT,
                    size=SIZE_NOTE, bold=True, margin=(80, 120))
        _write_cell(table.rows[ri].cells[1], val, val_width, HEX_ROW_BASE,
                    size=SIZE_NOTE, margin=(80, 120))
    return table


def build_cover(doc, title, meta, cfg):
    add_rich_paragraph(doc, cfg["confidential"], size=SIZE_NOTE,
                       color=COLOR_MUTED, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                       space_after=Pt(0))
    for _ in range(6):
        doc.add_paragraph()

    add_rich_paragraph(doc, title.replace(" 仕様書", ""), size=SIZE_TITLE,
                       color=COLOR_H1, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_rich_paragraph(doc, cfg["subtitle"], size=SIZE_SUBTITLE, color=COLOR_H2,
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_after=Pt(4))
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(30)
    _para_border(rule, "bottom", HEX_ACCENT_DARK, sz=6)

    add_rich_paragraph(doc, "対象ファイル情報", size=SIZE_NOTE, color=COLOR_H2,
                       bold=True, space_after=Pt(3))
    file_rows = [("対象ファイル", cfg["target_file"])]
    for key in ("ファイル種別", "用途", "VBA プロジェクト", "外部連携", "解析日"):
        if key in meta:
            file_rows.append((key, meta[key]))
    add_info_table(doc, file_rows)

    doc.add_paragraph()
    add_rich_paragraph(doc, "文書管理情報", size=SIZE_NOTE, color=COLOR_H2,
                       bold=True, space_after=Pt(3))
    mgmt_rows = [("文書番号", cfg["doc_no"]),
                 ("版数", f'第 {cfg["rev"]} 版'),
                 ("発行日", cfg["issue_date"]),
                 ("機密区分", cfg["confidential"])]
    if cfg["client"]:
        mgmt_rows.insert(0, ("提出先", cfg["client"]))
    if cfg["issuer"]:
        mgmt_rows.append(("発行元", cfg["issuer"]))
    add_info_table(doc, mgmt_rows)


def build_front_matter(doc, cfg):
    p = doc.add_paragraph("改訂履歴", style="Heading 1")
    _page_break_before(p)
    add_table(doc, ["版数", "発行日", "改訂内容", "作成", "確認"],
              [[f'{cfg["rev"]}', cfg["issue_date"], cfg["rev_note"], "", ""]])
    doc.add_paragraph()

    doc.add_paragraph("本書の取扱い", style="Heading 1")
    for text in cfg["handling"]:
        add_rich_paragraph(doc, f"・{text}", size=SIZE_BODY,
                           indent_left=Twips(180), space_after=Pt(3))
    doc.add_paragraph()

    doc.add_paragraph("本書の構成", style="Heading 1")
    add_rich_paragraph(
        doc,
        "本書は対象ブックの静的解析結果に基づき、シート構造・名前付き範囲・数式・"
        "VBA プロシージャ・DB 連携・データフロー・セキュリティ上の注意点を "
        "10 章構成で記述する。保守時に最初に確認すべき項目には ✓ を付す（判定基準は「凡例」章参照）。",
        size=SIZE_BODY)


def build_toc(doc, headings):
    """静的目次。TOC フィールドは埋め込まない（フィールド更新で頁がすべて 1 になるため）。"""
    p = doc.add_paragraph("目次", style="Heading 1")
    _page_break_before(p)

    for level, text in headings:
        if level == 2:
            add_rich_paragraph(doc, text, size=Pt(11), color=COLOR_H1, bold=True,
                               space_before=Pt(4), space_after=Pt(2))
        else:
            add_rich_paragraph(doc, text, size=Pt(10), color=COLOR_H2,
                               indent_left=Twips(360), space_after=Pt(1))


# ── 本文 ───────────────────────────────────────────────────

def build_body(doc, tokens):
    skip_title = True
    skip_meta_bq = True
    skip_toc_section = False

    for tok in tokens:
        kind = tok[0]

        if kind == "heading":
            level, text = tok[1], tok[2]
            if skip_title and level == 1:
                skip_title = False
                continue
            if level == 2 and plain_text(text).strip() == "目次":
                skip_toc_section = True   # md 側の目次は Word 目次で置き換える
                continue
            skip_toc_section = False

            if level == 2:
                p = doc.add_paragraph(plain_text(text), style="Heading 1")
                _page_break_before(p)
            elif level == 3:
                doc.add_paragraph(plain_text(text), style="Heading 2")
            elif level == 4:
                doc.add_paragraph(plain_text(text), style="Heading 3")
            else:
                doc.add_paragraph(plain_text(text), style="Heading 4")
            continue

        if skip_toc_section:
            continue

        if kind == "blockquote":
            if skip_meta_bq:
                skip_meta_bq = False
                continue
            text = " ".join(l for l in tok[1] if l.strip()).strip()
            if text:
                add_note(doc, text)

        elif kind == "paragraph":
            if re.match(r'^\*?以\s*上\*?\s*$', tok[1].strip()):
                continue
            add_rich_paragraph(doc, tok[1], size=SIZE_BODY)

        elif kind == "list_item":
            indent = Twips(240 + tok[2] * 200)
            add_rich_paragraph(doc, f"・{tok[1]}", size=SIZE_BODY,
                               indent_left=indent, space_after=Pt(2))

        elif kind == "table":
            add_table(doc, tok[1], tok[2])

        elif kind == "code":
            add_code_block(doc, list(tok[1]))


def build_closing(doc, cfg):
    doc.add_paragraph()
    p = add_rich_paragraph(doc, "以 上", size=SIZE_BODY, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           space_before=Pt(12))
    if cfg["contact"]:
        add_rich_paragraph(doc, f'本書に関するお問い合わせ先: {cfg["contact"]}',
                           size=SIZE_NOTE, color=COLOR_NOTE,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    return p


# ── メイン ────────────────────────────────────────────────

def build_docx(md_path, out_path, cfg):
    tokens = parse_md(md_path)
    title, meta = extract_meta(tokens)
    tokens = restructure(tokens)
    doc_name = plain_text(title) or Path(md_path).stem

    stem = doc_name.replace(" 仕様書", "")
    ext = re.match(r'(\.\w+)', meta.get("ファイル種別", ""))
    cfg.setdefault("target_file", stem + (ext.group(1) if ext else ".xlsm"))
    cfg.setdefault("doc_no", f"SPEC-{doc_name.replace(' 仕様書', '')}-{cfg['rev']}")

    doc = Document()
    setup_styles(doc)
    setup_page(doc, f'{doc_name}（{cfg["subtitle"]}）', cfg["confidential"],
               f'{cfg["doc_no"]} ／ 第 {cfg["rev"]} 版 ／ {cfg["issue_date"]}')

    headings = [(t[1], plain_text(t[2])) for t in tokens
                if t[0] == "heading" and t[1] in (2, 3)
                and plain_text(t[2]).strip() != "目次"]

    build_cover(doc, doc_name, meta, cfg)
    build_front_matter(doc, cfg)
    build_toc(doc, headings)
    build_body(doc, tokens)
    build_closing(doc, cfg)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"✅ 生成完了: {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Markdown 仕様書 → 納品用 docx")
    ap.add_argument("input", help="入力 .md")
    ap.add_argument("--out", help="出力 .docx（既定: 入力と同じ場所）")
    ap.add_argument("--subtitle", default="システム仕様書")
    ap.add_argument("--rev", default="1.0")
    ap.add_argument("--date", dest="issue_date", default=date.today().isoformat())
    ap.add_argument("--doc-no", dest="doc_no")
    ap.add_argument("--target-file", dest="target_file")
    ap.add_argument("--issuer", default="")
    ap.add_argument("--client", default="")
    ap.add_argument("--contact", default="")
    ap.add_argument("--confidential", default="機密文書 — 社外秘")
    ap.add_argument("--rev-note", dest="rev_note", default="初版発行")
    args = ap.parse_args()

    md_path = Path(args.input)
    out_path = Path(args.out) if args.out else md_path.with_suffix(".docx")

    cfg = {
        "subtitle": args.subtitle,
        "rev": args.rev,
        "issue_date": args.issue_date,
        "issuer": args.issuer,
        "client": args.client,
        "contact": args.contact,
        "confidential": args.confidential,
        "rev_note": args.rev_note,
        "handling": [
            "本書は対象ブックの保守を目的とした技術仕様書であり、記載内容には DB 接続情報等の機微な情報を含む。",
            "発行元の許諾なく、第三者への開示・複製・再配布を行わないこと。",
            "対象ブックの改修時は本書を併せて更新し、改訂履歴に記録すること。",
        ],
    }
    if args.doc_no:
        cfg["doc_no"] = args.doc_no
    if args.target_file:
        cfg["target_file"] = args.target_file

    build_docx(md_path, out_path, cfg)


if __name__ == "__main__":
    main()
